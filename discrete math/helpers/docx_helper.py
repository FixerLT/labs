from docx import Document
from wand.image import Image
from docx.shared import Mm


def create_docx():
    return Document()


def watermark_img(input_file, output_file):
    with Image(filename=input_file) as background:
        with Image(filename='/home/san/Documents/university/babakov/test/hui_watermark.png') as water_mark:
            background.watermark(image=water_mark, transparency=0.9)
        background.save(filename=output_file)


def add_step(document, header=None, image_path=None, watermark=False, comment=None, ender_page_break=False):
    if not ender_page_break:
        document.add_page_break()
    if header is not None:
        document.add_paragraph(header, style='Intense Quote')
    if image_path is not None:
        if watermark:
            watermark_img(image_path, '/home/san/Documents/university/babakov/test/watermark_tmp.png')
            document.add_picture('/home/san/Documents/university/babakov/test/watermark_tmp.png', width=Mm(140.5))
        else:
            document.add_picture(image_path)
    if comment is not None:
        document.add_paragraph(comment)
    if ender_page_break:
        document.add_page_break()


def merge_documents(documents):
    merged_document = Document()

    for index, file in enumerate(documents):

        if index > 0:
            merged_document.add_page_break()

        for element in file.element.body:
            merged_document.element.body.append(element)

    return merged_document


class PageReport:
    header = None
    image_path = None
    comment = None

    def __init__(self, header=None, image_path=None, comment=None):
        self.header = header
        self.image_path = image_path
        self.comment = comment

    def some(self):
        pass


# TODO: add method for merging two files
class Reporter:
    pages = None

    def save_report(self, path='/home/san/Documents/university/babakov/test/', report_name='report', watermark=True):
        document = Document()
        for page in self.pages:
            add_step(document, header=page.header, image_path=page.image_path,
                     watermark=watermark, comment=page.comment)
        document.save(path + report_name + '.docx')

    def add_page(self, page=None, header=None, image_path=None, comment=None):
        if page is None:
            page = PageReport()
        if header is not None:
            page.header = header
        if image_path is not None:
            page.image_path = image_path
        if comment is not None:
            page.comment = comment

        self.pages.append(page)

    def extend_last_comment(self, extra_text):
        if len(self.pages) == 0:
            self.add_page(comment='')
        if self.pages[-1].comment is None:
            self.pages[-1].comment = extra_text
        else:
            self.pages[-1].comment += extra_text

    def __init__(self):
        self.pages = []